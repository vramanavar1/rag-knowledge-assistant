"""End-to-end proof of the document lifecycle.

    python scripts/verify_lifecycle.py

Runs the whole add / modify / delete cycle against a throwaway copy of the
corpus in a temp directory, so the real index is untouched, and asserts the
behaviour that docs/ingestion-flow.md claims:

1. First ingest indexes everything.
2. Re-running with no changes parses nothing and embeds nothing.
3. A NEW document is indexed, and if it supersedes an existing one, that
   existing document is demoted *without being re-parsed or re-embedded*.
4. A MODIFIED document has its previous chunks purged before the new ones land,
   so no orphans survive.
5. A DELETED document leaves no chunks and stops being citable.

Exits non-zero if any assertion fails, so it can run in CI.
"""

from __future__ import annotations

import os
import shutil
import sys
import tempfile
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO_ROOT / "src"))

PASS, FAIL = "  PASS", "  FAIL"
_failures: list[str] = []


def check(label: str, condition: bool, detail: str = "") -> None:
    print(f"{PASS if condition else FAIL}  {label}{(' — ' + detail) if detail else ''}")
    if not condition:
        _failures.append(label)


def make_pricing_2027(path: Path) -> None:
    """A 2027 rate card that names the document it replaces."""
    try:
        import pymupdf as fitz
    except ImportError:  # PyMuPDF < 1.24.3
        import fitz

    body = (
        "OrbitSuite Pricing — 2027 Rate Card\n"
        "Northwind Traders, Inc. | Sales Operations | Effective: January 1, 2027 | "
        "Version 1.0 | Supersedes: 2026 Rate Card (v1.0)\n\n"
        "1. Overview\n"
        "This document summarizes standard list pricing for the OrbitSuite platform, "
        "effective for contracts signed on or after January 1, 2027. It replaces the "
        "2026 rate card and supersedes Pricing2026.pdf.\n\n"
        "2. Subscription Tiers\n"
        "Starter $35 per seat per month, minimum 5 seats.\n"
        "Professional $72 per seat per month, minimum 10 seats.\n"
        "Enterprise $119 per seat per month, minimum 25 seats.\n"
        "Enterprise Plus $159 per seat per month, minimum 100 seats.\n\n"
        "3. Notes\n"
        "Pricing reflects U.S. dollar list pricing effective January 1, 2027.\n"
    )
    doc = fitz.open()
    page = doc.new_page()
    page.insert_textbox(fitz.Rect(50, 50, 545, 780), body, fontsize=10, fontname="helv")
    doc.save(path)
    doc.close()


def modify_password_policy(path: Path) -> None:
    """Change a real value so the content hash moves."""
    import docx

    document = docx.Document(str(path))
    for paragraph in document.paragraphs:
        if "automatically locked after 5 consecutive failed" not in paragraph.text:
            continue
        # Word splits a sentence across runs at arbitrary points, so a
        # run-by-run replace usually matches nothing. Collapse the paragraph
        # into its first run and clear the rest.
        updated = paragraph.text.replace(
            "locked after 5 consecutive", "locked after 3 consecutive"
        )
        if paragraph.runs:
            paragraph.runs[0].text = updated
            for run in paragraph.runs[1:]:
                run.text = ""
        break

    document.add_paragraph("10. Review Cadence")
    document.add_paragraph(
        "This policy is reviewed every 6 months by the Security Governance Board."
    )
    document.save(str(path))


def main() -> int:
    from rag.config import get_settings
    from rag.ingest.manifest import Manifest
    from rag.ingest.sync import sync
    from rag.observability.tracing import configure_logging, set_correlation_id, start_trace
    from rag.providers.embeddings import get_embedding_provider
    from rag.store.local import LocalHybridStore

    configure_logging("ERROR", "text")

    workdir = Path(tempfile.mkdtemp(prefix="rag-lifecycle-"))
    source = workdir / "corpus"
    shutil.copytree(REPO_ROOT / "KnwoledgeBaseDocuments", source)

    os.environ["RAG_SOURCE_DIR"] = str(source)
    os.environ["RAG_DATA_DIR"] = str(workdir / "data")
    os.environ["RETRIEVER_BACKEND"] = "local"
    os.environ["RAG_PROFILE"] = "improved"
    settings = get_settings(refresh=True)

    def ingest():
        set_correlation_id(None)
        start_trace()
        embedder = get_embedding_provider(settings)
        store = LocalHybridStore(settings.index_path(), profile=settings.profile)
        store.load()
        report = sync(settings, store, embedder, source_dir=source)
        return report, store

    def manifest() -> Manifest:
        return Manifest(settings.manifest_path()).load()

    print(f"\nworkspace: {workdir}\n")

    # ---------------------------------------------------------------- 1
    print("1. First ingest")
    report, store = ingest()
    check("all documents indexed", len(report.new) == 11, f"{len(report.new)} new")
    check("chunks written", report.chunks_written > 100,
          f"{report.chunks_written} chunks")
    check("tables recovered", report.table_chunks >= 20,
          f"{report.table_chunks} table chunks")
    check("2025 rate card superseded by 2026",
          "Sales/Pricing2025.pdf" in report.superseded_now)
    baseline_chunks = report.total_chunks

    # ---------------------------------------------------------------- 2
    print("\n2. Re-ingest with no source changes")
    report, store = ingest()
    check("nothing reparsed",
          not report.new and not report.modified and not report.deleted,
          report.render().splitlines()[1].strip())
    check("no embedding API calls", report.embedding_calls == 0)
    check("chunk count unchanged", report.total_chunks == baseline_chunks,
          f"{report.total_chunks}")

    # ---------------------------------------------------------------- 3
    print("\n3. Add Pricing2027.pdf, modify PasswordPolicy.docx, delete VPNGuide.pdf")
    make_pricing_2027(source / "Sales" / "Pricing2027.pdf")
    modify_password_policy(source / "IT" / "PasswordPolicy.docx")
    (source / "IT" / "VPNGuide.pdf").unlink()

    before = manifest()
    vpn_chunk_ids = set(before.entries["IT/VPNGuide.pdf"].chunk_ids)
    old_password_chunk_ids = set(before.entries["IT/PasswordPolicy.docx"].chunk_ids)
    pricing2026_hash_before = before.entries["Sales/Pricing2026.pdf"].content_hash

    report, store = ingest()
    print("   " + report.render().replace("\n", "\n   "))

    check("classified correctly",
          report.new == ["Sales/Pricing2027.pdf"]
          and report.modified == ["IT/PasswordPolicy.docx"]
          and report.deleted == ["IT/VPNGuide.pdf"]
          and len(report.unchanged) == 9,
          f"new={report.new} mod={report.modified} del={report.deleted} "
          f"unchanged={len(report.unchanged)}")

    # ---- deletion
    remaining_vpn = [c for c in vpn_chunk_ids if store.get_chunk(c) is not None]
    check("deleted document left no chunks", not remaining_vpn,
          f"{len(remaining_vpn)} orphans")
    check("deleted document dropped from manifest",
          "IT/VPNGuide.pdf" not in manifest().entries)
    check("deleted document not retrievable",
          all(h.chunk.doc_id != "IT/VPNGuide.pdf"
              for h in store.search("How do I connect to the NorthLink VPN?",
                                    None, __import__("rag.store.base",
                                                     fromlist=["SearchFilters"])
                                    .SearchFilters(), 10, "keyword")))

    # ---- modification
    current_password_chunk_ids = set(
        manifest().entries["IT/PasswordPolicy.docx"].chunk_ids
    )
    stale = [
        c for c in (old_password_chunk_ids - current_password_chunk_ids)
        if store.get_chunk(c) is not None
    ]
    check("modified document left no stale chunks", not stale,
          f"{len(stale)} orphans")
    lockout = [
        c for c in store.chunks_for_doc("IT/PasswordPolicy.docx")
        if "3 consecutive failed" in c.text
    ]
    check("modified content is what is now indexed", bool(lockout))
    check("new section indexed",
          any("Review Cadence" in c.section_path or "Review Cadence" in c.text
              for c in store.chunks_for_doc("IT/PasswordPolicy.docx")))

    # ---- supersession without re-parsing
    entries = manifest().entries
    check("2027 rate card is current", entries["Sales/Pricing2027.pdf"].meta["is_current"])
    check("2026 rate card demoted by the new arrival",
          entries["Sales/Pricing2026.pdf"].meta["is_current"] is False,
          f"superseded_by={entries['Sales/Pricing2026.pdf'].meta['superseded_by']}")
    check("2026 rate card was NOT re-parsed (hash unchanged)",
          entries["Sales/Pricing2026.pdf"].content_hash == pricing2026_hash_before)
    check("2026 rate card was in the 'unchanged' bucket",
          "Sales/Pricing2026.pdf" in report.unchanged)
    check("demotion reached the index, not just the manifest",
          all(not c.is_current for c in store.chunks_for_doc("Sales/Pricing2026.pdf")))
    check("2025 rate card still superseded",
          entries["Sales/Pricing2025.pdf"].meta["is_current"] is False)

    # ---------------------------------------------------------------- 4
    print("\n4. Remove Pricing2027.pdf again")
    (source / "Sales" / "Pricing2027.pdf").unlink()
    report, store = ingest()
    entries = manifest().entries
    check("2026 rate card reinstated as current",
          entries["Sales/Pricing2026.pdf"].meta["is_current"] is True,
          "reconciliation is reversible")
    check("reinstatement reported", "Sales/Pricing2026.pdf" in report.reinstated)
    check("reinstatement reached the index",
          all(c.is_current for c in store.chunks_for_doc("Sales/Pricing2026.pdf")))

    shutil.rmtree(workdir, ignore_errors=True)

    print()
    if _failures:
        print(f"FAILED: {len(_failures)} check(s): {', '.join(_failures)}")
        return 1
    print("All lifecycle checks passed.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
