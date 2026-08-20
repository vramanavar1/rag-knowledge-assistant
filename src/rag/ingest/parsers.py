"""Structure-preserving parsers for PDF, DOCX and XLSX.

Why this file is bigger than "extract the text"
-----------------------------------------------
Every parser here produces TWO renderings of the same document:

* ``blocks``      -- headings, paragraphs and tables in true reading order,
                    with tables rendered as Markdown.  Used by the *improved*
                    profile.
* ``naive_text``  -- what you get from the obvious one-liner in each library.
                    Used by the *baseline* profile.

The difference is not cosmetic, and it is the measured root cause of the
assignment's Scenario 1 ("correct document, wrong chunk"):

* PDF   -- ``page.get_text()`` emits a table column-by-column, so
           "Client meals / $100/person / Director" arrives as three unrelated
           runs of text and the row association is destroyed.  ``find_tables()``
           recovers the grid.
* DOCX  -- ``doc.paragraphs`` skips tables entirely and ``doc.tables`` returns
           them detached, so the hotel rate caps ($350 / $250 / $180) end up
           orphaned from the "4. Hotel Accommodations" heading they belong to.
           Iterating ``body.iterchildren()`` restores the interleaving.
* XLSX  -- raw cell values ignore number formats, so a 15% discount indexes as
           "0.15".  We apply the workbook's number formats.

All three parsers are stdlib + PyMuPDF + python-docx only; the XLSX reader is
written against the OOXML zip directly so that openpyxl is not required.
"""

from __future__ import annotations

import html
import re
import zipfile
from dataclasses import dataclass, field
from pathlib import Path

from rag.observability.tracing import get_logger

log = get_logger(__name__)

BULLET_CHARS = {"•", "●", "▪", "◦", "-", "–", "*", "o"}

# "1. Purpose", "2.1 Annual / Paid Time Off (PTO)", "10. Governing Law"
HEADING_RE = re.compile(r"^(\d+(?:\.\d+)*)[.)]?\s+(\S.{0,90})$")

# The metadata strapline every document in this corpus carries, e.g.
# "Northwind Traders, Inc. | Finance Department | Effective: February 1, 2026 | Version 5.1 | ..."
HEADER_HINT_RE = re.compile(
    r"(effective|plan year|last updated|template version|version)\s*[:\s]", re.I
)


# --------------------------------------------------------------------------
# Result types
# --------------------------------------------------------------------------


@dataclass
class Block:
    kind: str                 # "title" | "heading" | "text" | "table"
    text: str
    level: int = 0
    page: int | None = None


@dataclass
class ParsedDoc:
    title: str
    header_line: str
    blocks: list[Block] = field(default_factory=list)
    page_count: int | None = None
    naive_text: str = ""

    @property
    def table_count(self) -> int:
        return sum(1 for b in self.blocks if b.kind == "table")


# --------------------------------------------------------------------------
# Shared helpers
# --------------------------------------------------------------------------


def _clean_cell(value: object) -> str:
    if value is None:
        return ""
    text = str(value).replace("\r", " ").replace("\n", " ").strip()
    text = re.sub(r"\s+", " ", text)
    return text.replace("|", r"\|")


def rows_to_markdown(rows: list[list[object]]) -> str:
    """Render a grid as a Markdown table, dropping empty rows and columns."""
    grid = [[_clean_cell(c) for c in row] for row in rows]
    grid = [r for r in grid if any(c for c in r)]
    if not grid:
        return ""

    width = max(len(r) for r in grid)
    grid = [r + [""] * (width - len(r)) for r in grid]

    keep = [i for i in range(width) if any(r[i] for r in grid)]
    if not keep:
        return ""
    grid = [[r[i] for i in keep] for r in grid]

    if len(grid) == 1:
        return " | ".join(grid[0])

    header, body = grid[0], grid[1:]
    # A blank header cell reads badly in Markdown; give it a placeholder.
    header = [h or f"col{i + 1}" for i, h in enumerate(header)]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join("---" for _ in header) + " |",
    ]
    lines += ["| " + " | ".join(r) + " |" for r in body]
    return "\n".join(lines)


def rows_to_key_values(rows: list[list[object]]) -> str:
    """Render a two-column grid as a bullet list.

    Spreadsheet blocks like "Billing Term | Annual (prepaid)" are key/value
    pairs, not tables; forcing them into Markdown tables invents a header row
    that misleads the model.
    """
    out = []
    for row in rows:
        cells = [_clean_cell(c) for c in row if _clean_cell(c)]
        if not cells:
            continue
        if len(cells) == 1:
            out.append(f"- {cells[0]}")
        else:
            key = cells[0].rstrip(":").strip()
            out.append(f"- {key}: {' | '.join(cells[1:])}")
    return "\n".join(out)


def _is_bullet_marker(text: str) -> bool:
    """A standalone bullet glyph, laid out as its own line by the PDF.

    Matches the known glyphs plus any 1-2 character line with no alphanumerics,
    which covers the Symbol/Wingdings private-use bullets that Word-generated
    PDFs often carry.
    """
    stripped = text.strip()
    if not stripped:
        return False
    if stripped in BULLET_CHARS:
        return True
    return len(stripped) <= 2 and not any(ch.isalnum() for ch in stripped)


def _is_heading(line: str) -> tuple[bool, int, str]:
    m = HEADING_RE.match(line.strip())
    if not m:
        return False, 0, ""
    number, rest = m.group(1), m.group(2)
    # "12. of the Agreement" style false positives: a heading should not end
    # mid-sentence, and legal prose numbers clauses inline.
    if rest.endswith((",", ";")) or len(rest.split()) > 12:
        return False, 0, ""
    level = number.count(".") + 1
    return True, level, f"{number}. {rest}".strip()


def _find_header_line(candidates: list[str]) -> str:
    for line in candidates[:8]:
        if HEADER_HINT_RE.search(line) and ("|" in line or "  " in line):
            return line.strip()
    for line in candidates[:8]:
        if HEADER_HINT_RE.search(line):
            return line.strip()
    return ""


# --------------------------------------------------------------------------
# PDF
# --------------------------------------------------------------------------


def _detect_boilerplate(page_lines: list[list[str]]) -> set[str]:
    """Lines repeated across most pages are running headers/footers.

    Short, non-alphanumeric lines are excluded on purpose: a bare "•" sits in
    the first or last three lines of nearly every page, and treating it as a
    running header silently deletes every bullet in the document.
    """
    if len(page_lines) < 2:
        return set()
    counts: dict[str, int] = {}
    for lines in page_lines:
        for line in set(lines[:3] + lines[-3:]):
            stripped = line.strip()
            if len(stripped) < 4 or len(stripped) > 90:
                continue
            if not any(ch.isalnum() for ch in stripped):
                continue
            counts[stripped] = counts.get(stripped, 0) + 1
    threshold = max(2, int(len(page_lines) * 0.6))
    boiler = {t for t, c in counts.items() if c >= threshold}
    boiler |= {t for t in counts if re.fullmatch(r"Page \d+( of \d+)?", t, re.I)}
    return boiler


def parse_pdf(path: Path) -> ParsedDoc:
    # PyMuPDF renamed its module from `fitz` to `pymupdf`; 1.28.2 prints a
    # deprecation warning on every `import fitz`, which lands in the middle of
    # normal ingest output. Prefer the new name, fall back for older versions.
    try:
        import pymupdf as fitz
    except ImportError:  # PyMuPDF < 1.24.3
        import fitz

    doc = fitz.open(path)
    naive_text = "\n".join(page.get_text() for page in doc)

    # Pass 1: collect per-page lines so running headers/footers can be spotted.
    per_page_lines: list[list[str]] = []
    for page in doc:
        per_page_lines.append([ln for ln in page.get_text().splitlines() if ln.strip()])
    boilerplate = _detect_boilerplate(per_page_lines)

    blocks: list[Block] = []
    first_lines: list[str] = []

    for page_no, page in enumerate(doc, start=1):
        try:
            found = page.find_tables()
            tables = list(found.tables)
        except Exception as exc:  # pragma: no cover - defensive
            log.warning("table detection failed", page=page_no, error=str(exc))
            tables = []

        table_rects = [fitz.Rect(t.bbox) for t in tables]

        # Ordered items on this page: (y, x, kind, payload)
        items: list[tuple[float, float, str, object]] = []

        for tbl, rect in zip(tables, table_rects):
            try:
                markdown = rows_to_markdown(tbl.extract())
            except Exception as exc:  # pragma: no cover - defensive
                log.warning("table extract failed", page=page_no, error=str(exc))
                continue
            if markdown:
                items.append((rect.y0, rect.x0, "table", markdown))

        text_dict = page.get_text("dict")
        for block in text_dict.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                text = "".join(span.get("text", "") for span in line.get("spans", []))
                if not text.strip():
                    continue
                bbox = fitz.Rect(line["bbox"])
                # Skip anything that lives inside a detected table: it is
                # already represented, with its row structure intact.
                if any(rect.intersects(bbox) and
                       (rect & bbox).get_area() > 0.5 * bbox.get_area()
                       for rect in table_rects):
                    continue
                if text.strip() in boilerplate:
                    continue
                size = max((span.get("size", 0.0) for span in line.get("spans", [])),
                           default=0.0)
                items.append((bbox.y0, bbox.x0, "line", (text.strip(), size, bbox.y1)))

        items.sort(key=lambda it: (round(it[0], 1), it[1]))

        # Merge consecutive lines into paragraphs; headings and bullets break them.
        buffer: list[str] = []
        pending_bullet = False
        prev_y1: float | None = None

        def flush() -> None:
            nonlocal buffer
            if buffer:
                blocks.append(Block("text", " ".join(buffer).strip(), page=page_no))
                buffer = []

        for y0, _x0, kind, payload in items:
            if kind == "table":
                flush()
                blocks.append(Block("table", str(payload), page=page_no))
                prev_y1 = None
                continue

            text, size, y1 = payload  # type: ignore[misc]
            if page_no == 1 and len(first_lines) < 10:
                first_lines.append(text)

            if _is_bullet_marker(text):
                flush()
                pending_bullet = True
                prev_y1 = y1
                continue

            is_head, level, normalised = _is_heading(text)
            if is_head:
                flush()
                blocks.append(Block("heading", normalised, level=level, page=page_no))
                pending_bullet = False
                prev_y1 = y1
                continue

            if pending_bullet:
                flush()
                buffer.append(f"- {text}")
                pending_bullet = False
            else:
                # A large vertical gap means a new paragraph.
                if prev_y1 is not None and y0 - prev_y1 > 6.0:
                    flush()
                buffer.append(text)
            prev_y1 = y1

        flush()

    title = ""
    for line in first_lines:
        if line.strip() and not HEADER_HINT_RE.search(line):
            title = line.strip()
            break
    header_line = _find_header_line(first_lines)

    # The cover lines land in one merged paragraph because they sit close
    # together; split the title back out so the front matter reads correctly.
    if title and blocks and blocks[0].kind == "text" and blocks[0].text.startswith(title):
        remainder = blocks[0].text[len(title):].strip()
        blocks[0] = Block("title", title, page=1)
        if remainder:
            blocks.insert(1, Block("text", remainder, page=1))

    return ParsedDoc(
        title=title or path.stem,
        header_line=header_line,
        blocks=blocks,
        page_count=doc.page_count,
        naive_text=naive_text,
    )


# --------------------------------------------------------------------------
# DOCX
# --------------------------------------------------------------------------


def parse_docx(path: Path) -> ParsedDoc:
    import docx
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    document = docx.Document(str(path))

    # The naive rendering the baseline profile uses: paragraphs first, then
    # every table bolted on at the end, detached from its section.
    naive_parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            naive_parts.append(" | ".join(c.text.strip() for c in row.cells))
    naive_text = "\n".join(naive_parts)

    blocks: list[Block] = []
    title = ""
    first_lines: list[str] = []
    body = document.element.body

    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]

        if tag == "tbl":
            table = Table(child, document)
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            markdown = rows_to_markdown(rows)
            if markdown:
                blocks.append(Block("table", markdown))
            continue

        if tag != "p":
            continue

        paragraph = Paragraph(child, document)
        text = paragraph.text.strip()
        if not text:
            continue

        style = paragraph.style.name if paragraph.style is not None else ""
        if len(first_lines) < 10:
            first_lines.append(text)

        if style == "Title" and not title:
            title = text
            blocks.append(Block("title", text))
            continue

        is_head, level, normalised = _is_heading(text)
        if style.startswith("Heading") or is_head:
            try:
                style_level = int(style.split()[-1])
            except (ValueError, IndexError):
                style_level = level or 1
            blocks.append(Block("heading", normalised or text, level=style_level))
            continue

        if style == "List Paragraph":
            blocks.append(Block("text", f"- {text}"))
            continue

        blocks.append(Block("text", text))

    return ParsedDoc(
        title=title or path.stem,
        header_line=_find_header_line(first_lines),
        blocks=blocks,
        page_count=None,
        naive_text=naive_text,
    )


# --------------------------------------------------------------------------
# XLSX  (OOXML read directly -- no openpyxl dependency)
# --------------------------------------------------------------------------

_BUILTIN_NUM_FORMATS = {
    9: "0%",
    10: "0.00%",
    14: "date",
    44: '"$"#,##0.00',
    164: "General",
}

_CELL_RE = re.compile(r"<c\s([^>]*?)(?:/>|>(.*?)</c>)", re.S)
_ROW_RE = re.compile(r"<row[^>]*>(.*?)</row>", re.S)


def _column_index(ref: str) -> int:
    """'C7' -> 2 (zero-based column)."""
    letters = "".join(ch for ch in ref if ch.isalpha())
    idx = 0
    for ch in letters:
        idx = idx * 26 + (ord(ch.upper()) - 64)
    return max(idx - 1, 0)


def _format_number(raw: str, format_code: str) -> str:
    try:
        value = float(raw)
    except ValueError:
        return raw

    if "%" in format_code:
        decimals = 0
        if "." in format_code:
            decimals = len(format_code.split(".")[1].split("%")[0].rstrip("%"))
        return f"{value * 100:.{decimals}f}%"

    if "$" in format_code:
        decimals = 2 if ".00" in format_code else 0
        return f"${value:,.{decimals}f}"

    if value == int(value):
        return str(int(value))
    return f"{value:g}"


def _read_sheet_rows(sheet_xml: str, shared: list[str],
                     style_formats: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for row_xml in _ROW_RE.findall(sheet_xml):
        cells: dict[int, str] = {}
        for attrs, body in _CELL_RE.findall(row_xml):
            body = body or ""
            ref = re.search(r'r="([A-Z]+\d+)"', attrs)
            col = _column_index(ref.group(1)) if ref else len(cells)

            cell_type = re.search(r't="(\w+)"', attrs)
            cell_type = cell_type.group(1) if cell_type else "n"
            style_idx = re.search(r's="(\d+)"', attrs)
            format_code = ""
            if style_idx:
                i = int(style_idx.group(1))
                if 0 <= i < len(style_formats):
                    format_code = style_formats[i]

            if cell_type == "inlineStr":
                m = re.search(r"<t[^>]*>(.*?)</t>", body, re.S)
                value = html.unescape(re.sub(r"<[^>]+>", "", m.group(1))) if m else ""
            else:
                m = re.search(r"<v>(.*?)</v>", body, re.S)
                if not m:
                    continue
                raw = html.unescape(m.group(1))
                if cell_type == "s":
                    value = shared[int(raw)] if raw.isdigit() and int(raw) < len(shared) else raw
                elif cell_type in ("str", "e"):
                    value = raw
                else:
                    value = _format_number(raw, format_code)
            if value.strip():
                cells[col] = value.strip()

        if cells:
            width = max(cells) + 1
            rows.append([cells.get(i, "") for i in range(width)])
        else:
            rows.append([])
    return rows


def parse_xlsx(path: Path) -> ParsedDoc:
    zf = zipfile.ZipFile(path)
    names = zf.namelist()

    shared: list[str] = []
    if "xl/sharedStrings.xml" in names:
        xml = zf.read("xl/sharedStrings.xml").decode("utf-8", "replace")
        for item in re.findall(r"<si>(.*?)</si>", xml, re.S):
            shared.append(html.unescape(re.sub(r"<[^>]+>", "", item)))

    # style index -> number format code
    style_formats: list[str] = []
    if "xl/styles.xml" in names:
        styles_xml = zf.read("xl/styles.xml").decode("utf-8", "replace")
        custom = {
            int(fid): code
            for fid, code in re.findall(
                r'<numFmt numFmtId="(\d+)" formatCode="([^"]*)"', styles_xml
            )
        }
        cell_xfs = re.search(r"<cellXfs[^>]*>(.*?)</cellXfs>", styles_xml, re.S)
        if cell_xfs:
            for xf in re.findall(r"<xf\s[^>]*?/?>", cell_xfs.group(1)):
                fid = re.search(r'numFmtId="(\d+)"', xf)
                fid_int = int(fid.group(1)) if fid else 0
                style_formats.append(
                    custom.get(fid_int, _BUILTIN_NUM_FORMATS.get(fid_int, ""))
                )

    sheet_names: list[str] = []
    if "xl/workbook.xml" in names:
        wb = zf.read("xl/workbook.xml").decode("utf-8", "replace")
        sheet_names = [html.unescape(n) for n in re.findall(r'<sheet name="([^"]+)"', wb)]

    blocks: list[Block] = []
    naive_parts: list[str] = []
    sheet_files = sorted(n for n in names if re.fullmatch(r"xl/worksheets/sheet\d+\.xml", n))

    title = ""
    header_line = ""

    for idx, sheet_file in enumerate(sheet_files):
        sheet_name = sheet_names[idx] if idx < len(sheet_names) else f"Sheet{idx + 1}"
        rows = _read_sheet_rows(zf.read(sheet_file).decode("utf-8", "replace"),
                                shared, style_formats)

        blocks.append(Block("heading", f"{idx + 1}. {sheet_name}", level=1))
        naive_parts.append(f"[{sheet_name}]")

        # Group consecutive rows by how many populated columns they have; a
        # change in shape means a new logical block on the sheet.
        group: list[list[str]] = []
        group_width = -1

        def flush_group() -> None:
            nonlocal group, group_width
            if not group:
                return
            populated = [[c for c in row if c.strip()] for row in group]
            width = max((len(r) for r in populated), default=0)
            if width <= 1:
                for row in populated:
                    if row:
                        blocks.append(Block("text", row[0]))
            elif width == 2:
                text = rows_to_key_values(group)
                if text:
                    blocks.append(Block("table", text))
            else:
                markdown = rows_to_markdown(group)
                if markdown:
                    blocks.append(Block("table", markdown))
            group = []
            group_width = -1

        for row in rows:
            populated = len([c for c in row if c.strip()])
            if populated == 0:
                flush_group()
                continue
            naive_parts.append(" | ".join(row))
            if group_width != -1 and populated != group_width:
                flush_group()
            group_width = populated
            group.append(row)
        flush_group()

        if idx == 0:
            for block in blocks:
                if block.kind == "text" and not title:
                    title = block.text
                if HEADER_HINT_RE.search(block.text) and not header_line:
                    header_line = block.text

    return ParsedDoc(
        title=title or path.stem,
        header_line=header_line,
        blocks=blocks,
        page_count=None,
        naive_text="\n".join(naive_parts),
    )


# --------------------------------------------------------------------------
# Dispatch
# --------------------------------------------------------------------------

_PARSERS = {".pdf": parse_pdf, ".docx": parse_docx, ".xlsx": parse_xlsx}

SUPPORTED_EXTENSIONS = frozenset(_PARSERS)


def parse_document(path: Path) -> ParsedDoc:
    ext = path.suffix.lower()
    parser = _PARSERS.get(ext)
    if parser is None:
        raise ValueError(f"unsupported document type: {ext}")
    parsed = parser(path)
    log.debug(
        "parsed document",
        path=str(path),
        blocks=len(parsed.blocks),
        tables=parsed.table_count,
    )
    return parsed
